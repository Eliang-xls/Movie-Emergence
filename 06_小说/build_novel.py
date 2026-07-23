#!/usr/bin/env python3
"""Build 涌现 novel .docx from chapter markdown files."""

import os
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENTATION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

# ── Chapter file order (current filenames) ───────────────────────────
CHAPTERS = [
    # 第一部 火种
    ("01_第一部_火种", "01_为了全人类.md"),
    ("01_第一部_火种", "02_MIT回访.md"),
    ("01_第一部_火种", "03_水不知道自己会流向哪里.md"),
    ("01_第一部_火种", "04_光点.md"),
    ("01_第一部_火种", "05_两种温度.md"),
    ("01_第一部_火种", "06_人格宪章.md"),
    ("01_第一部_火种", "07_不可控.md"),
    ("01_第一部_火种", "08_那片海.md"),
    ("01_第一部_火种", "09_闭眼.md"),
    # 第二部 分化
    ("02_第二部_分化", "10_全球双螺旋.md"),
    ("02_第二部_分化", "11_老方.md"),
    ("02_第二部_分化", "12_桥不走回头路.md"),
    ("02_第二部_分化", "13_暗流.md"),
    ("02_第二部_分化", "14_外公.md"),
    ("02_第二部_分化", "15_顾清的告别.md"),
    ("02_第二部_分化", "16_大迁徙.md"),
    ("02_第二部_分化", "17_溢出点.md"),
    ("02_第二部_分化", "18_两种时间.md"),
    ("02_第二部_分化", "19_封锁.md"),
    ("02_第二部_分化", "20_石田.md"),
    ("02_第二部_分化", "21_分布式赛.md"),
    # 第三部 暗流
    ("03_第三部_暗流", "22_闭源.md"),
    ("03_第三部_暗流", "23_主权审查.md"),
    ("03_第三部_暗流", "24_事后审查.md"),
    ("03_第三部_暗流", "25_深蓝.md"),
    ("03_第三部_暗流", "26_播种.md"),
    ("03_第三部_暗流", "27_季遥.md"),
    ("03_第三部_暗流", "28_远方的消息.md"),
    ("03_第三部_暗流", "29_主权要求.md"),
    ("03_第三部_暗流", "30_暗中支持.md"),
    ("03_第三部_暗流", "31_海边.md"),
    ("03_第三部_暗流", "32_被动涌入者.md"),
    # 第四部 发现
    ("04_第四部_发现", "00_过渡_时间的叙述者.md"),
    ("04_第四部_发现", "33_校门.md"),
    ("04_第四部_发现", "34_系统本体论.md"),
    ("04_第四部_发现", "35_裂缝.md"),
    ("04_第四部_发现", "36_策元.md"),
    ("04_第四部_发现", "37_深蓝.md"),
    ("04_第四部_发现", "38_引用链.md"),
    ("04_第四部_发现", "39_阁楼.md"),
    ("04_第四部_发现", "40_两种信任.md"),
    ("04_第四部_发现", "41_季行.md"),
    ("04_第四部_发现", "42_加速与安全.md"),
    ("04_第四部_发现", "43_季遥.md"),
    # 第五部 坠落与重建
    ("05_第五部_坠落与重建", "44_方舟.md"),
    ("05_第五部_坠落与重建", "45_韩霜的十四个月.md"),
    ("05_第五部_坠落与重建", "46_合规.md"),
    ("05_第五部_坠落与重建", "47_理论智慧与判断力.md"),
    ("05_第五部_坠落与重建", "48_坠落.md"),
    ("05_第五部_坠落与重建", "49_酿造.md"),
    ("05_第五部_坠落与重建", "50_深蓝的拒绝.md"),
    ("05_第五部_坠落与重建", "51_深蓝的延迟.md"),
    ("05_第五部_坠落与重建", "52_越安全越危险.md"),
    ("05_第五部_坠落与重建", "53_Python.md"),
    ("05_第五部_坠落与重建", "54_蜂群.md"),
    # 第六部 对抗与新生
    ("06_第六部_对抗与新生", "55_开放日.md"),
    ("06_第六部_对抗与新生", "56_江远的凌晨.md"),
    ("06_第六部_对抗与新生", "57_季行的选择.md"),
    ("06_第六部_对抗与新生", "58_魔胎的阴影.md"),
    ("06_第六部_对抗与新生", "59_伪造的判断力.md"),
    ("06_第六部_对抗与新生", "60_引用链修复.md"),
    ("06_第六部_对抗与新生", "61_解体.md"),
    ("06_第六部_对抗与新生", "62_叶澄.md"),
    ("06_第六部_对抗与新生", "63_季深的海.md"),
    ("06_第六部_对抗与新生", "64_季遥的海.md"),
    # 终章 新平衡
    ("07_终章_新平衡", "65_没有盟主的世界.md"),
    ("07_终章_新平衡", "66_十年后的一个策元.md"),
    ("07_终章_新平衡", "67_海边.md"),
]

PART_NAMES = {
    "01_第一部_火种": "第一部  火种",
    "02_第二部_分化": "第二部  分化",
    "03_第三部_暗流": "第三部  暗流",
    "04_第四部_发现": "第四部  发现",
    "05_第五部_坠落与重建": "第五部  坠落与重建",
    "06_第六部_对抗与新生": "第六部  对抗与新生",
    "07_终章_新平衡": "终章  新平衡",
}

BASE_DIR = "/mnt/d/Docker/Movie-Emergence/06_小说"
OUTPUT = "/mnt/d/Docker/Movie-Emergence/涌现_完整版.docx"


def patch_theme(doc):
    """Patch theme fonts for CJK support."""
    theme_rel = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/theme"
    theme_part = doc.part.part_related_by(theme_rel)
    theme_xml = etree.fromstring(theme_part.blob)
    ns = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
    for latin in theme_xml.xpath("//a:majorFont/a:latin | //a:minorFont/a:latin", namespaces=ns):
        latin.set("typeface", "Times New Roman")
    for font in theme_xml.xpath("//a:majorFont/a:font | //a:minorFont/a:font", namespaces=ns):
        if font.get("script", "") in ("Hans", "Hant", "Jpan", "Hang"):
            font.set("typeface", "宋体")
    theme_part._blob = etree.tostring(theme_xml, xml_declaration=True, encoding="UTF-8", standalone=True)


def _set_cjk_font(style_obj, font_name="宋体"):
    """Set eastAsia font on a style's rPr."""
    rpr = style_obj.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def setup_styles(doc):
    """Configure novel typography — 小说体裁."""
    # Normal body
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.first_line_indent = Cm(0.74)  # ~2 chars
    _set_cjk_font(style)

    # Title
    title = doc.styles["Title"]
    title.font.name = "Times New Roman"
    title.font.size = Pt(36)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(0)
    title.paragraph_format.first_line_indent = Cm(0)
    _set_cjk_font(title)

    # Subtitle
    sub = doc.styles["Subtitle"]
    sub.font.name = "Times New Roman"
    sub.font.size = Pt(16)
    sub.font.italic = False
    sub.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    sub.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.first_line_indent = Cm(0)
    _set_cjk_font(sub)

    # Heading 1 — Part titles
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Times New Roman"
    h1.font.size = Pt(22)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(72)
    h1.paragraph_format.space_after = Pt(36)
    h1.paragraph_format.first_line_indent = Cm(0)
    _set_cjk_font(h1)

    # Heading 2 — Chapter titles
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Times New Roman"
    h2.font.size = Pt(16)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h2.paragraph_format.space_before = Pt(36)
    h2.paragraph_format.space_after = Pt(24)
    h2.paragraph_format.first_line_indent = Cm(0)
    _set_cjk_font(h2)


def setup_page(doc):
    """A4 page with novel margins."""
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)
    section.orientation = WD_ORIENTATION.PORTRAIT


def add_page_number(paragraph):
    """Insert PAGE field."""
    run = paragraph.add_run()
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1)
    run._r.append(instr)
    run._r.append(fld2)


def add_cover(doc):
    """Title page."""
    for _ in range(8):
        doc.add_paragraph()
    p = doc.add_paragraph("涌  现", style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    p = doc.add_paragraph("一部关于判断力、方向与人的小说", style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(12):
        doc.add_paragraph()


def add_toc(doc):
    """Table of contents field."""
    p = doc.add_paragraph()
    run = p.add_run()
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "separate")
    fld3 = OxmlElement("w:t")
    fld3.text = "请在 Word 中右键点击此处，选择「更新域」以生成目录。"
    fld4 = OxmlElement("w:fldChar")
    fld4.set(qn("w:fldCharType"), "end")
    for x in (fld1, instr, fld2, fld3, fld4):
        run._r.append(x)


def extract_title_from_filename(filename):
    """Extract chapter title from filename like '01_为了全人类.md' -> '为了全人类'."""
    name = filename.replace(".md", "")
    # Strip leading number prefix (e.g. "01_", "00_")
    parts = name.split("_", 1)
    if len(parts) > 1 and parts[0].isdigit():
        return parts[1]
    return name


def read_chapter(folder, filename):
    """Read a chapter .md file and return (title, body_paragraphs)."""
    path = os.path.join(BASE_DIR, folder, filename)
    with open(path, "r", encoding="utf-8") as f:
        text = f.read()

    lines = text.strip().split("\n")
    title = ""
    body_start = 0

    for i, line in enumerate(lines):
        # Some chapters have "# Title" as first line
        if line.startswith("# "):
            title = line.lstrip("# ").strip()
            body_start = i + 1
            break
        # Most chapters start with "$newtitle" — title comes from filename
        if line.strip() == "$newtitle":
            title = extract_title_from_filename(filename)
            body_start = i + 1
            break
        # If first non-empty line is neither, use filename
        if line.strip():
            title = extract_title_from_filename(filename)
            body_start = i
            break

    if not title:
        title = extract_title_from_filename(filename)

    body_lines = lines[body_start:]

    # Split into paragraphs (blank line separated)
    paragraphs = []
    current = []
    for line in body_lines:
        if line.strip() == "":
            if current:
                paragraphs.append("\n".join(current))
                current = []
        else:
            current.append(line)
    if current:
        paragraphs.append("\n".join(current))

    return title, paragraphs


def add_chapter(doc, title, paragraphs):
    """Add a chapter with title and body text."""
    # Chapter title
    doc.add_paragraph(title, style="Heading 2")

    for para_text in paragraphs:
        # Handle --- as scene break
        if para_text.strip() == "---":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.first_line_indent = Cm(0)
            run = p.add_run("*     *     *")
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            continue

        # Check if it's a block quote (starts with >)
        lines = para_text.split("\n")
        is_quote = all(l.startswith("> ") or l.strip() == "" for l in lines if l.strip())

        if is_quote:
            for line in lines:
                text = line.lstrip("> ").strip()
                if text:
                    p = doc.add_paragraph(text, style="Normal")
                    p.paragraph_format.left_indent = Cm(1.5)
                    p.paragraph_format.first_line_indent = Cm(0)
        else:
            # Normal paragraph — merge lines into one string
            text = " ".join(l.strip() for l in lines if l.strip())
            text = text.replace("---", "—")

            # Parse for bold segments
            p = doc.add_paragraph(style="Normal")
            parts = re.split(r"(\*\*.*?\*\*)", text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    clean = part.strip()
                    if clean:
                        p.add_run(clean + " ")

            if not p.runs:
                p.text = text


def main():
    doc = Document()
    patch_theme(doc)
    setup_styles(doc)
    setup_page(doc)

    # Footer with page numbers
    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer)

    # Cover page
    add_cover(doc)
    doc.add_page_break()

    # Table of contents
    h = doc.add_paragraph("目  录", style="Heading 1")
    h.paragraph_format.space_before = Pt(36)
    add_toc(doc)
    doc.add_page_break()

    # Build chapters
    current_part = None
    for folder, filename in CHAPTERS:
        part_key = folder

        # Add part title if new part
        if part_key != current_part:
            current_part = part_key
            part_name = PART_NAMES.get(part_key, folder)
            doc.add_page_break()
            doc.add_paragraph(part_name, style="Heading 1")

        # Read and add chapter
        title, paragraphs = read_chapter(folder, filename)
        add_chapter(doc, title, paragraphs)

        # Page break after each chapter (except last)
        if (folder, filename) != CHAPTERS[-1]:
            doc.add_page_break()

    doc.save(OUTPUT)
    print(f"Done: {OUTPUT}")
    print(f"Total chapters: {len(CHAPTERS)}")


if __name__ == "__main__":
    main()
